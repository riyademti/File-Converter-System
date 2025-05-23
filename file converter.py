# Multifeature Bangla Converter App - Complete Working Version

import tkinter as tk
from tkinter import ttk, filedialog, messagebox
import sqlite3
import os
import hashlib
from PIL import Image
import webbrowser
try:
    from moviepy.editor import VideoFileClip
except ModuleNotFoundError:
    VideoFileClip = None
from docx import Document
from fpdf import FPDF
import fitz

# ─── CONFIGURATION ─────────────────────────────────────────────────────────────
APP_TITLE = "Bangla Converter"
WINDOW_SIZE = "880x600"

# Theme settings: background, foreground, panel, button, button_hover
THEMES = {
    "dark": {
        "bg": "#1a1a1a",
        "fg": "#f9fafb",
        "panel": "#2b2b2b",
        "button": "#3a84ff",
        "button_hover": "#336fcc"
    },
    "light": {
        "bg": "#f9fafb",
        "fg": "#111827",
        "panel": "#ffffff",
        "button": "#007acc",
        "button_hover": "#006bb3"
    }
}
current_theme = "dark"
current_lang = "bn"

# Localization strings
LANGUAGES = {
    "bn": {
        "title": "বাংলা কনভার্টার",
        "login": "লগইন",
        "register": "রেজিস্টার",
        "username": "ইউজারনেম",
        "password": "পাসওয়ার্ড",
        "reset": "পাসওয়ার্ড রিসেট",
        "theme": "থিম",
        "history": "হিস্টোরি",
        "converter_options": "কনভার্টার অপশন",
        "logout": "লগআউট",
        "invalid_credentials": "ইউজারনেম অথবা পাসওয়ার্ড ভুল",
        "username_exists": "ইউজারনেম ইতিমধ্যে আছে",
        "registration_success": "নিবন্ধন সফল হয়েছে"
    },
    "en": {
        "title": "Bangla Converter",
        "login": "Login",
        "register": "Register",
        "username": "Username",
        "password": "Password",
        "reset": "Reset Password",
        "theme": "Theme",
        "history": "History",
        "converter_options": "Converter Options",
        "logout": "Logout",
        "invalid_credentials": "Invalid username or password",
        "username_exists": "This username already exists",
        "registration_success": "Registration successful"
    }
}

DB_PATH = 'users.db'

# ─── DATABASE SETUP ────────────────────────────────────────────────────────────
def init_db():
    with sqlite3.connect(DB_PATH) as conn:
        c = conn.cursor()
        c.execute('''CREATE TABLE IF NOT EXISTS users (
                        username TEXT PRIMARY KEY,
                        password TEXT
                    )''')
        c.execute('''CREATE TABLE IF NOT EXISTS recent_files (
                        file_path TEXT,
                        action TEXT,
                        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                    )''')
        conn.commit()

# ─── UTILITY FUNCTIONS ────────────────────────────────────────────────────────
def hash_password(pw: str) -> str:
    return hashlib.sha256(pw.encode()).hexdigest()

def check_credentials(username: str, password: str) -> bool:
    with sqlite3.connect(DB_PATH) as conn:
        c = conn.cursor()
        c.execute("SELECT 1 FROM users WHERE username=? AND password=?", (username, hash_password(password)))
        return c.fetchone() is not None

def register_user(username: str, password: str) -> bool:
    try:
        with sqlite3.connect(DB_PATH) as conn:
            c = conn.cursor()
            c.execute("INSERT INTO users (username, password) VALUES (?,?)", (username, hash_password(password)))
            conn.commit()
        return True
    except sqlite3.IntegrityError:
        return False

def log_recent(file_path: str, action: str):
    with sqlite3.connect(DB_PATH) as conn:
        c = conn.cursor()
        c.execute("INSERT INTO recent_files (file_path, action) VALUES (?,?)", (file_path, action))
        conn.commit()

def get_recent_files(limit: int = 8):
    with sqlite3.connect(DB_PATH) as conn:
        c = conn.cursor()
        c.execute("SELECT file_path, action FROM recent_files ORDER BY created_at DESC LIMIT ?", (limit,))
        return c.fetchall()

# ─── CONVERTER FUNCTIONS ──────────────────────────────────────────────────────
def convert_docx_to_pdf():
    path = filedialog.askopenfilename(title="Select DOCX File", filetypes=[("Word Document", "*.docx")])
    if not path: return
    try:
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        doc = Document(path)
        for para in doc.paragraphs:
            pdf.multi_cell(0, 8, para.text)
        out = os.path.splitext(path)[0] + '.pdf'
        pdf.output(out)
        log_recent(out, "DOCX→PDF")
        messagebox.showinfo("Success", f"Saved: {out}")
        webbrowser.open(out)
    except Exception as e:
        messagebox.showerror("Error", str(e))


def convert_pdf_to_docx():
    path = filedialog.askopenfilename(title="Select PDF File", filetypes=[("PDF File", "*.pdf")])
    if not path: return
    try:
        doc_pdf = fitz.open(path)
        text = "".join(page.get_text() for page in doc_pdf)
        out_doc = Document()
        for line in text.splitlines():
            out_doc.add_paragraph(line)
        out = os.path.splitext(path)[0] + '.docx'
        out_doc.save(out)
        log_recent(out, "PDF→DOCX")
        messagebox.showinfo("Success", f"Saved: {out}")
        webbrowser.open(out)
    except Exception as e:
        messagebox.showerror("Error", str(e))


def convert_jpg_to_png():
    path = filedialog.askopenfilename(title="Select JPG File", filetypes=[("JPEG Image", "*.jpg")])
    if not path: return
    try:
        out = os.path.splitext(path)[0] + '.png'
        Image.open(path).save(out)
        log_recent(out, "JPG→PNG")
        messagebox.showinfo("Success", f"Saved: {out}")
        webbrowser.open(out)
    except Exception as e:
        messagebox.showerror("Error", str(e))


def convert_png_to_jpg():
    path = filedialog.askopenfilename(title="Select PNG File", filetypes=[("PNG Image", "*.png")])
    if not path: return
    try:
        out = os.path.splitext(path)[0] + '.jpg'
        Image.open(path).convert("RGB").save(out)
        log_recent(out, "PNG→JPG")
        messagebox.showinfo("Success", f"Saved: {out}")
        webbrowser.open(out)
    except Exception as e:
        messagebox.showerror("Error", str(e))


def convert_video_to_mp3():
    if VideoFileClip is None:
        messagebox.showerror("Error", "Install moviepy: pip install moviepy")
        return
    path = filedialog.askopenfilename(title="Select Video File", filetypes=[("Video Files", ".mp4;.avi;*.mov")])
    if not path: return
    try:
        clip = VideoFileClip(path)
        out = os.path.splitext(path)[0] + '.mp3'
        clip.audio.write_audiofile(out)
        clip.close()
        log_recent(out, "Video→MP3")
        messagebox.showinfo("Success", f"Saved: {out}")
        webbrowser.open(out)
    except Exception as e:
        messagebox.showerror("Error", str(e))

# ─── UI HELPERS ───────────────────────────────────────────────────────────────
def on_enter(btn):
    btn['bg'] = THEMES[current_theme]['button_hover']

def on_leave(btn):
    btn['bg'] = THEMES[current_theme]['button']

# ─── SCREENS ─────────────────────────────────────────────────────────────────
def login_screen():
    for w in root.winfo_children():
        w.destroy()
    th = THEMES[current_theme]
    txt = LANGUAGES[current_lang]
    root.configure(bg=th['bg'])

    frm = tk.Frame(root, bg=th['bg'])
    frm.place(relx=0.5, rely=0.5, anchor='center')

    tk.Label(frm, text=txt['username'], font=('Helvetica',12), bg=th['bg'], fg=th['fg']).grid(row=0, column=0, pady=(0,5), sticky='w')
    user_entry = tk.Entry(frm, width=30)
    user_entry.grid(row=1, column=0, pady=(0,10))

    tk.Label(frm, text=txt['password'], font=('Helvetica',12), bg=th['bg'], fg=th['fg']).grid(row=2, column=0, pady=(10,5), sticky='w')
    pass_entry = tk.Entry(frm, width=30, show='*')
    pass_entry.grid(row=3, column=0, pady=(0,15))

    def try_login():
        if check_credentials(user_entry.get(), pass_entry.get()):
            open_main_window()
        else:
            messagebox.showerror("Error", txt['invalid_credentials'])

    def try_register():
        if register_user(user_entry.get(), pass_entry.get()):
            messagebox.showinfo("Success", txt['registration_success'])
        else:
            messagebox.showwarning("Warning", txt['username_exists'])

    btn_login = tk.Button(frm, text=txt['login'], font=('Helvetica',11), bg=th['button'], fg='white', command=try_login)
    btn_login.grid(row=4, column=0, sticky='we', pady=(0,5))
    btn_login.bind('<Enter>', lambda e: on_enter(btn_login))
    btn_login.bind('<Leave>', lambda e: on_leave(btn_login))

    btn_reg = tk.Button(frm, text=txt['register'], font=('Helvetica',11), bg=th['button'], fg='white', command=try_register)
    btn_reg.grid(row=5, column=0, sticky='we', pady=(0,5))
    btn_reg.bind('<Enter>', lambda e: on_enter(btn_reg))
    btn_reg.bind('<Leave>', lambda e: on_leave(btn_reg))

    btn_reset = tk.Button(frm, text=txt['reset'], font=('Helvetica',11), bg=th['button'], fg='white', command=lambda: messagebox.showinfo("Info","Coming soon"))
    btn_reset.grid(row=6, column=0, sticky='we', pady=(0,5))

    btn_theme = tk.Button(frm, text=txt['theme'], font=('Helvetica',11), bg=th['button'], fg='white', command=toggle_theme)
    btn_theme.grid(row=7, column=0, sticky='we', pady=(0,5))

    btn_lang = tk.Button(frm, text='🌐', font=('Helvetica',11), bg=th['button'], fg='white', command=toggle_language)
    btn_lang.grid(row=8, column=0, sticky='we', pady=(0,5))


def open_main_window():
    for w in root.winfo_children():
        w.destroy()
    th = THEMES[current_theme]
    txt = LANGUAGES[current_lang]
    root.configure(bg=th['bg'])

    sidebar = tk.Frame(root, bg=th['panel'], width=200)
    sidebar.pack(side='left', fill='y')

    tk.Label(sidebar, text=txt['converter_options'], font=('Helvetica',14,'bold'), bg=th['panel'], fg=th['button']).pack(pady=(10,20))
    btns = [
        ("📄 DOCX→PDF", convert_docx_to_pdf),
        ("📄 PDF→DOCX", convert_pdf_to_docx),
        ("🖼 JPG→PNG", convert_jpg_to_png),
        ("🖼 PNG→JPG", convert_png_to_jpg),
        ("🎞 Video→MP3", convert_video_to_mp3)
    ]
    for text, cmd in btns:
        btn = tk.Button(sidebar, text=text, font=('Helvetica',11), bg=th['button'], fg='white', command=cmd)
        btn.pack(fill='x', pady=5, padx=10)
        btn.bind('<Enter>', lambda e, b=btn: on_enter(b))
        btn.bind('<Leave>', lambda e, b=btn: on_leave(b))

    tk.Label(sidebar, text=txt['history'], font=('Helvetica',12), bg=th['panel'], fg=th['fg']).pack(pady=(30,5))
    for path, action in get_recent_files():
        tk.Label(sidebar, text=f"{action}\n{os.path.basename(path)}", font=('Helvetica',8), bg=th['panel'], fg=th['fg'], anchor='w', justify='left', wraplength=180).pack(fill='x', padx=10, pady=2)

    tk.Button(sidebar, text=txt['logout'], font=('Helvetica',11), bg='#e53e3e', fg='white', command=login_screen).pack(side='bottom', fill='x', padx=10, pady=10)

# ─── THEME & LANGUAGE TOGGLES ─────────────────────────────────────────────────
def toggle_theme():
    global current_theme
    current_theme = 'light' if current_theme=='dark' else 'dark'
    login_screen()

def toggle_language():
    global current_lang
    current_lang = 'en' if current_lang=='bn' else 'bn'
    login_screen()

# ─── MAIN ─────────────────────────────────────────────────────────────────────
if __name__ == '__main__':
    init_db()
    root = tk.Tk()
    root.title(APP_TITLE)
    root.geometry(WINDOW_SIZE)
    login_screen()
    root.mainloop()